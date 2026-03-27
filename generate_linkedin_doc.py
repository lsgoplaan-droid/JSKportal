from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
import os

output_dir = r"d:\Claude Projects\JSK AIAutomation and Digital Solutions"
img_dir = os.path.join(output_dir, "linkedin_images")
os.makedirs(img_dir, exist_ok=True)

def create_banner(filename, title, subtitle, colors, icons_text):
    w, h = 1200, 400
    img = Image.new('RGB', (w, h), colors[0])
    draw = ImageDraw.Draw(img)

    for i in range(w):
        r1, g1, b1 = colors[0]
        r2, g2, b2 = colors[1]
        ratio = i / w
        r = int(r1 + (r2 - r1) * ratio)
        g = int(g1 + (g2 - g1) * ratio)
        b = int(b1 + (b2 - b1) * ratio)
        draw.line([(i, 0), (i, h)], fill=(r, g, b))

    draw.ellipse([w-250, -100, w+50, 200], fill=(*colors[1],))
    draw.ellipse([-80, h-200, 120, h+20], fill=(*colors[0],))

    for x in range(0, w, 60):
        draw.line([(x, 0), (x, h)], fill=(*colors[1],), width=1)
    for y in range(0, h, 60):
        draw.line([(0, y), (w, y)], fill=(*colors[1],), width=1)

    draw.rectangle([60, 80, w-60, h-80], fill=(10, 10, 20, 200))

    try:
        font_title = ImageFont.truetype("arial.ttf", 42)
        font_sub = ImageFont.truetype("arial.ttf", 22)
        font_icons = ImageFont.truetype("arial.ttf", 18)
    except:
        font_title = ImageFont.load_default()
        font_sub = ImageFont.load_default()
        font_icons = ImageFont.load_default()

    draw.text((100, 110), title, fill=(255, 255, 255), font=font_title)
    draw.text((100, 175), subtitle, fill=(200, 200, 220), font=font_sub)

    tag_x = 100
    for tag in icons_text:
        bbox = draw.textbbox((0, 0), tag, font=font_icons)
        tw = bbox[2] - bbox[0]
        draw.rounded_rectangle([tag_x, 240, tag_x + tw + 24, 275], radius=6, fill=(255, 255, 255, 40))
        draw.text((tag_x + 12, 245), tag, fill=(255, 255, 255), font=font_icons)
        tag_x += tw + 40

    img.save(os.path.join(img_dir, filename))
    print(f"Created {filename}")

# Hero
create_banner(
    "hero_agentic.png",
    "From Temple Verses to Trade Finance",
    "How a Personal Project Led Me Down the Agentic AI Rabbit Hole",
    [(14, 14, 34), (99, 102, 241)],
    ["Agentic AI", "Claude Code", "Autonomous Reasoning", "Human-in-the-Loop"]
)

# Chapter 1: Kavya
create_banner(
    "kavya_poetry.png",
    "Kavya \u2014 Poetry Translator",
    "It Started With a Simple Question: Can AI Understand These Verses?",
    [(20, 184, 166), (99, 102, 241)],
    ["Claude AI", "Expo", "TypeScript", "Supabase"]
)

# Chapter 2: Trading Scanner
create_banner(
    "trading_scanner.png",
    "Trading Opportunity Scanner",
    "What If This Thing Could Watch Markets While I Sleep?",
    [(99, 102, 241), (139, 92, 246)],
    ["FastAPI", "NumPy", "Pandas", "WebSocket"]
)

# Chapter 3: LC Checker
create_banner(
    "lc_checker.png",
    "LC Compliance Checker",
    "The Real Test: Could It Solve a Problem I Spent Years Doing Manually?",
    [(5, 150, 105), (37, 99, 235)],
    ["FastAPI", "React", "OCR", "AI Agent"]
)

# Workflow diagram
def create_workflow():
    w, h = 1200, 350
    img = Image.new('RGB', (w, h), (14, 14, 34))
    draw = ImageDraw.Draw(img)

    try:
        font = ImageFont.truetype("arial.ttf", 16)
        font_sm = ImageFont.truetype("arial.ttf", 13)
        font_title = ImageFont.truetype("arial.ttf", 24)
    except:
        font = ImageFont.load_default()
        font_sm = ImageFont.load_default()
        font_title = ImageFont.load_default()

    draw.text((w//2 - 180, 20), "Agentic AI Development Workflow", fill=(255, 255, 255), font=font_title)

    steps = [
        ("Goal\nSpecification", "Human defines\noutcome", (0, 232, 130)),
        ("Autonomous\nPlanning", "Agent decomposes\nproblem", (99, 102, 241)),
        ("Multi-Step\nExecution", "Agent writes &\niterates code", (168, 85, 247)),
        ("Self-\nEvaluation", "Agent reviews\nown output", (236, 72, 153)),
        ("Human\nReview", "Domain expert\nvalidates", (0, 232, 130)),
    ]

    box_w, box_h = 180, 120
    start_x = 40
    y = 120

    for i, (title, desc, color) in enumerate(steps):
        x = start_x + i * 230
        draw.rounded_rectangle([x, y, x+box_w, y+box_h], radius=12, fill=(30, 30, 60), outline=color, width=2)
        draw.text((x + 15, y + 15), title, fill=color, font=font)
        draw.text((x + 15, y + 65), desc, fill=(180, 180, 200), font=font_sm)

        if i < len(steps) - 1:
            arrow_start = x + box_w + 5
            arrow_end = x + box_w + 45
            arrow_y = y + box_h // 2
            draw.line([(arrow_start, arrow_y), (arrow_end, arrow_y)], fill=(100, 100, 140), width=2)
            draw.polygon([(arrow_end, arrow_y-6), (arrow_end+10, arrow_y), (arrow_end, arrow_y+6)], fill=(100, 100, 140))

    draw.text((w//2 - 100, y + box_h + 30), "Iterate until goal is met", fill=(100, 100, 140), font=font_sm)
    img.save(os.path.join(img_dir, "workflow.png"))
    print("Created workflow.png")

create_workflow()

# --- Build Word Document ---
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(30, 30, 30)
style.paragraph_format.space_after = Pt(8)

def add_img(path):
    doc.add_picture(os.path.join(img_dir, path), width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading_colored(text, level, color):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_body(text):
    doc.add_paragraph(text)

def add_italic(text, color=RGBColor(100, 100, 120)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.italic = True
    run.font.color.rgb = color

# ==================== TITLE ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("From Temple Verses to Trade Finance")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(20, 20, 40)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("How a Personal Project Led Me Down the Agentic AI Rabbit Hole")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(99, 102, 241)

doc.add_paragraph("")
add_img("hero_agentic.png")
doc.add_paragraph("")

# ==================== INTRO ====================
add_body(
    "This story doesn't start with a grand plan to explore agentic AI. It doesn't start with a business case "
    "or a technology thesis. It starts with something far simpler."
)
add_body(
    "It starts with a prayer."
)

# ==================== CHAPTER 1 ====================
doc.add_paragraph("")
add_img("kavya_poetry.png")
add_heading_colored("Chapter 1: The Verses That Started Everything", 1, RGBColor(20, 184, 166))

add_body(
    "I grew up listening to my grandmother recite Tamil devotional poetry — the Divya Prabandham, "
    "Thirukkural, verses from the 108 Divyadesam temples. Beautiful, ancient words that I could hear "
    "but never fully understand. The classical Tamil in these poems is dense — agglutinative morphology, "
    "layered meanings, centuries of literary tradition compressed into a few lines."
)
add_body(
    "I'd always wanted a way to truly understand what each verse meant. Not just a rough translation — "
    "the literal meaning of each line, the poetic interpretation, the historical context, the scholarly "
    "significance. The kind of explanation you'd get from a patient Tamil literature professor, not Google Translate."
)
add_body(
    "So I tried something. I opened Claude Code — an agentic AI coding tool — and described what I wanted: "
    "a mobile app where I could type, photograph, or speak a Tamil verse, and get back a rich, structured "
    "translation with line-by-line breakdown, poetic interpretation in English, modern Tamil rendering, "
    "Hindi translation, and scholarly metadata."
)
add_body(
    "I expected to spend weeks building it. The agent had a working prototype in a single session."
)
add_body(
    "It scaffolded an Expo React Native project with file-based routing. It set up a Claude API integration "
    "with a 65-line system prompt I refined to handle the nuances of Sangam-era secular poetry differently "
    "from Bhakti devotional literature. It built four input methods — keyboard, camera with base64 image "
    "pipeline for vision model inference, speech-to-text, and clipboard. It configured expo-secure-store for "
    "API key management on native and localStorage fallback on web. It wired up Supabase for analytics and "
    "AsyncStorage with a 100-entry LIFO buffer for translation history."
)
add_body(
    "I didn't write a single line of code. I described what I wanted, tested it with real verses from the "
    "Divya Prabandham, told the agent where the translations felt wrong, and watched it iterate. The cultural "
    "knowledge — knowing that a particular verse uses a specific raga reference, or that the sandhi compound "
    "in a Sanskrit shloka carries a double meaning — that was mine. The implementation was the agent's."
)
add_body(
    "I named the app Kavya. It works. I use it regularly. And it changed how I thought about what's possible."
)
add_italic("Stack: Expo Router, TypeScript strict mode, Claude API (claude-haiku-4-5-20251001), Supabase, GitHub Actions CI/CD.")

# ==================== CHAPTER 2 ====================
doc.add_paragraph("")
add_img("trading_scanner.png")
add_heading_colored("Chapter 2: What If It Could Watch Markets While I Sleep?", 1, RGBColor(99, 102, 241))

add_body(
    "Kavya planted a seed. If an AI agent could navigate the complexity of classical Tamil poetry — "
    "morphology, literary traditions, multi-script output — what could it do with structured numerical data?"
)
add_body(
    "I've traded markets for years. The tedious part was never the decision-making — it was the scanning. "
    "Watching charts across US equities, India NSE, options chains, and FX pairs. Running the same technical "
    "checks over and over. SMA/EMA crossovers. MACD divergence. RSI momentum shifts. Bollinger Band squeezes. "
    "It's repetitive, time-consuming, and the moment you look away, you miss something."
)
add_body(
    "So I described the problem to the agent: build me a digital analyst that monitors multi-asset markets "
    "autonomously, runs technical analysis across a configurable set of indicators, scores opportunities "
    "using a weighted model, and — critically — backtests every signal against historical data before "
    "surfacing it."
)
add_body(
    "The agent designed an observe-orient-decide pipeline. Data ingestion through WebSocket feeds. A NumPy "
    "and Pandas-powered computation layer for time-series analysis. A multi-factor scoring engine that weights "
    "signals by reliability and market regime. A backtesting module that validates strategies before they "
    "reach the dashboard."
)
add_body(
    "What struck me wasn't just that it worked. It was how the agent reasoned about the architecture. When I "
    "said \"I don't want false signals,\" it didn't just add a threshold filter. It proposed a confirmation "
    "mechanism — requiring convergence across multiple independent indicators before scoring an opportunity "
    "as actionable. That's the kind of thinking I'd expect from a quantitative analyst, not a coding tool."
)
add_body(
    "The agent was operating as an autonomous reasoning system — decomposing my goal into sub-problems, "
    "proposing solutions for each, self-evaluating its own architecture, and iterating. This wasn't "
    "autocomplete. This was agentic AI in action: multi-step planning with tool use, goal-directed "
    "execution, and self-correction."
)
add_body(
    "Now I was genuinely impressed. And I started thinking about a much harder problem."
)
add_italic("Stack: FastAPI, React dashboard, NumPy/Pandas for quantitative analysis, WebSocket feeds.")

# ==================== CHAPTER 3 ====================
doc.add_paragraph("")
add_img("lc_checker.png")
add_heading_colored("Chapter 3: The Real Test — A Problem I Spent Years Solving Manually", 1, RGBColor(5, 150, 105))

add_body(
    "In my years in commodity trading and trade finance, I've sat through countless Letter of Credit reviews. "
    "If you've worked in trade ops, you know the pain. A bank issues an LC. Documents come in — invoices, "
    "bills of lading, certificates of origin, inspection reports. Every single field has to be checked against "
    "UCP600 rules and the LC terms. Is the beneficiary name exactly right? Does the shipping date fall within "
    "the validity period? Does the goods description match letter-for-letter?"
)
add_body(
    "One missed discrepancy can delay payment by weeks. One overlooked clause can cost real money. Experienced "
    "trade ops analysts spend hours on a single set of documents, and they still miss things when they're "
    "fatigued."
)
add_body(
    "This was the problem I knew better than any other. And if the agent could handle Tamil poetry and "
    "market analysis, I wanted to see if it could handle this."
)
add_body(
    "I described the domain in detail. UCP600 articles. The types of discrepancies that matter — and the "
    "ones that don't. How banks actually review documents in practice. The difference between a critical "
    "discrepancy that stops payment and a minor one that gets waived."
)
add_body(
    "The agent built something that genuinely surprised me."
)
add_body(
    "It designed a multi-step agentic pipeline: PDF ingestion with a dual-path extraction system — native "
    "text extraction for digital documents, Tesseract OCR fallback for scanned copies. An entity recognition "
    "layer that identifies and extracts key fields (beneficiary, applicant, amounts, dates, goods description, "
    "shipping terms). A fuzzy matching engine using Levenshtein distance and token-set ratio for field "
    "comparison — because in trade documents, a slight variation in company name spelling shouldn't trigger "
    "a false discrepancy. And a severity classification system that grades findings as critical, major, or "
    "minor, with confidence scores."
)
add_body(
    "The architecture wasn't what I would have designed. It was better. The agent proposed a confidence "
    "scoring mechanism I hadn't considered — assigning a reliability score to each extracted field based on "
    "OCR quality, then adjusting the compliance verdict accordingly. A scanned document with poor image "
    "quality gets flagged for human review on uncertain fields, while a clean digital PDF runs fully "
    "autonomously."
)
add_body(
    "That's the moment I understood what agentic AI really means. It's not about generating code faster. "
    "It's about an autonomous system that can reason through a complex problem, propose a solution architecture, "
    "execute it, evaluate its own output, and refine — all within a goal-directed loop that I'm guiding but "
    "not micromanaging."
)
add_italic("Stack: FastAPI, React, Python OCR pipeline, fuzzy matching engine, structured JSON output with confidence scoring.")

# ==================== WORKFLOW ====================
doc.add_paragraph("")
add_heading_colored("What I Actually Learned About Human-Agent Collaboration", 1, RGBColor(99, 102, 241))
add_img("workflow.png")
doc.add_paragraph("")

add_heading_colored("The agent gets better when you describe problems, not solutions.", 2, RGBColor(50, 50, 70))
add_body(
    "My early instinct was to over-specify. \"Use this library. Structure the API this way.\" The output was "
    "mediocre — I was constraining the agent's reasoning chain. The breakthrough came when I switched to "
    "outcome-based goal specification: \"Build a compliance pipeline that a junior trade ops analyst could "
    "trust without escalating.\" The agent's autonomous planning produced better architectures than my "
    "prescriptive instructions every single time."
)

add_heading_colored("Context accumulation is the real multiplier.", 2, RGBColor(50, 50, 70))
add_body(
    "The magic wasn't in any single prompt. It was in the accumulated context over a long session. By hour "
    "two of building Kavya, I could say \"add camera input with the same error handling pattern\" and the "
    "agent knew exactly which service layer, which error boundary, and which UI component to update. It had "
    "built a mental model of my project — design tokens, data flow, deployment pipeline, my preferences."
)
add_body(
    "This is why \"prompt engineering\" is a misleading term. What actually matters is session engineering — "
    "the ability to build a productive, context-rich working relationship with an AI agent over an extended "
    "collaboration. The value compounds with shared context, exactly like working with a colleague who's "
    "been on the project for months."
)

add_heading_colored("The agent makes confident mistakes. That's where you earn your keep.", 2, RGBColor(50, 50, 70))
add_body(
    "The agent once referenced a CSS class that didn't exist in the codebase. It occasionally hallucinated "
    "API parameters. On the LC checker, it initially proposed a string-matching approach for goods "
    "descriptions that anyone with trade finance experience would know doesn't work — because real documents "
    "use abbreviations, alternate spellings, and formatting variations that defeat exact matching."
)
add_body(
    "I caught it. I corrected it. And the agent immediately adapted — that's when it proposed the fuzzy "
    "matching approach that ended up being far more robust."
)
add_body(
    "The human-in-the-loop isn't a safety net. It's the evaluation function. The agent handles the "
    "execution chain. The human provides the judgment that the agent can't learn from training data alone — "
    "the kind that comes from years of actually doing the work."
)

# ==================== CONCLUSION ====================
doc.add_paragraph("")
add_heading_colored("Where This Leaves Me", 1, RGBColor(99, 102, 241))

add_body(
    "I started this journey trying to understand my grandmother's prayers. I ended up rethinking how "
    "software gets built."
)
add_body(
    "These three projects taught me that agentic AI doesn't replace domain expertise. It amplifies it. "
    "My 20 years in financial services and trading didn't become less valuable — they became dramatically "
    "more leveraged. The problems I can now tackle as a single person with an AI agent would have required "
    "a cross-functional team and months of development time just two years ago."
)
add_body(
    "The shift isn't from human to machine. It's from implementation to orchestration. From writing code "
    "to directing autonomous reasoning systems that write code. From being a specialist who executes to "
    "being a generalist who guides."
)
add_body(
    "The professionals who will thrive in this era aren't the fastest coders. They're the ones with enough "
    "domain depth to specify the right goals, enough technical fluency to evaluate an agent's reasoning chain, "
    "and enough real-world experience to catch the confident mistakes that only a human would notice."
)

p = doc.add_paragraph()
run = p.add_run(
    "And sometimes, the most important thing a human brings is knowing which question to ask in the first "
    "place — even when that question is as simple as \"what does this verse actually mean?\""
)
run.font.bold = True
run.font.size = Pt(12)

doc.add_paragraph("")
doc.add_paragraph("")

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "I'm exploring opportunities where I can apply this human-agent collaboration approach — "
    "combining deep domain expertise in financial services and trading with AI-augmented development.\n"
    "If that resonates with what you're building, I'd love to connect."
)
run.font.italic = True
run.font.color.rgb = RGBColor(99, 102, 241)
run.font.size = Pt(11)

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("What's been your experience with agentic AI? I'd love to hear your story in the comments.")
run.font.italic = True
run.font.color.rgb = RGBColor(120, 120, 140)

output_path = os.path.join(output_dir, "LinkedIn_Article_Agentic_AI_v2.docx")
doc.save(output_path)
print(f"\nDocument saved to: {output_path}")
