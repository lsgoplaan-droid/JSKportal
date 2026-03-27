from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageDraw, ImageFont
import os

output_dir = r"d:\Claude Projects\JSK AIAutomation and Digital Solutions"
img_dir = os.path.join(output_dir, "linkedin_images")
os.makedirs(img_dir, exist_ok=True)

def create_banner(filename, title, subtitle, colors, icons_text):
    w, h = 1200, 400
    img = Image.new('RGB', (w, h), colors[0])
    draw = ImageDraw.Draw(img)

    for i in range(w):
        r1, g1, b1 = colors[0]
        r2, g2, b2 = colors[1]
        ratio = i / w
        r = int(r1 + (r2 - r1) * ratio)
        g = int(g1 + (g2 - g1) * ratio)
        b = int(b1 + (b2 - b1) * ratio)
        draw.line([(i, 0), (i, h)], fill=(r, g, b))

    draw.ellipse([w-250, -100, w+50, 200], fill=(*colors[1],))
    draw.ellipse([-80, h-200, 120, h+20], fill=(*colors[0],))

    for x in range(0, w, 60):
        draw.line([(x, 0), (x, h)], fill=(*colors[1],), width=1)
    for y in range(0, h, 60):
        draw.line([(0, y), (w, y)], fill=(*colors[1],), width=1)

    draw.rectangle([60, 80, w-60, h-80], fill=(10, 10, 20, 200))

    try:
        font_title = ImageFont.truetype("arial.ttf", 38)
        font_sub = ImageFont.truetype("arial.ttf", 20)
        font_icons = ImageFont.truetype("arial.ttf", 16)
    except:
        font_title = ImageFont.load_default()
        font_sub = ImageFont.load_default()
        font_icons = ImageFont.load_default()

    draw.text((100, 110), title, fill=(255, 255, 255), font=font_title)
    draw.text((100, 170), subtitle, fill=(200, 200, 220), font=font_sub)

    tag_x = 100
    for tag in icons_text:
        bbox = draw.textbbox((0, 0), tag, font=font_icons)
        tw = bbox[2] - bbox[0]
        draw.rounded_rectangle([tag_x, 230, tag_x + tw + 24, 262], radius=6, fill=(255, 255, 255, 40))
        draw.text((tag_x + 12, 235), tag, fill=(255, 255, 255), font=font_icons)
        tag_x += tw + 40

    img.save(os.path.join(img_dir, filename))
    print(f"Created {filename}")

# Hero banner
create_banner(
    "company_hero.png",
    "Stop Worrying About LLMs and RAG",
    "Here's What AI Automation Actually Looks Like for Small Businesses",
    [(14, 14, 34), (99, 102, 241)],
    ["AI Automation", "MSMEs", "No Jargon", "Practical Solutions"]
)

# Section banners
create_banner(
    "company_workflows.png",
    "5 Workflows You Can Automate Today",
    "No PhD Required. No Developer Needed.",
    [(5, 150, 105), (99, 102, 241)],
    ["Invoices", "Customer Queries", "Lead Follow-ups", "Reports"]
)

create_banner(
    "company_litmus.png",
    "The 30-Second Litmus Test",
    "Is Your Business Ready for AI Automation?",
    [(99, 102, 241), (139, 92, 246)],
    ["Same Steps Daily", "30+ Minutes", "No Creative Judgment"]
)

# --- Build Word Document ---
doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor(30, 30, 30)
style.paragraph_format.space_after = Pt(8)

def add_img(path):
    doc.add_picture(os.path.join(img_dir, path), width=Inches(6))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_heading_colored(text, level, color):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_body(text):
    doc.add_paragraph(text)

def add_bold_body(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    return p

def add_italic(text, color=RGBColor(100, 100, 120)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.italic = True
    run.font.color.rgb = color

# ==================== FORMATTING NOTE ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("LINKEDIN ARTICLE — Publish from JSK AI Company Page")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(99, 102, 241)
run.font.bold = True

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Copy the text below into LinkedIn's article editor. Upload the banner images separately.")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(120, 120, 140)
run.font.italic = True

doc.add_paragraph("")
doc.add_paragraph("—" * 40)
doc.add_paragraph("")

# ==================== TITLE ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Stop Worrying About LLMs and RAG")
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = RGBColor(20, 20, 40)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Here's What AI Automation Actually Looks Like for Small Businesses")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(99, 102, 241)

doc.add_paragraph("")
add_img("company_hero.png")
doc.add_paragraph("")

# ==================== INTRO ====================
add_body(
    "Every week, a new AI acronym trends on LinkedIn. LLMs. RAG. SLMs. Fine-tuning. "
    "Agentic frameworks."
)
add_body(
    "If you're running a small business, this noise doesn't help you. It confuses you. "
    "And that confusion is costing you money — not because you're missing out on AI, "
    "but because the hype is stopping you from seeing what's already possible."
)
add_body("Let's cut through it.")

# ==================== SECTION 1 ====================
doc.add_paragraph("")
add_heading_colored(
    "What AI Automation Actually Looks Like for an MSME",
    1, RGBColor(5, 150, 105)
)

add_body(
    "Forget the buzzwords. Here's what AI automation means in practice for a business "
    "doing $500K–$5M in revenue:"
)

doc.add_paragraph("")

add_bold_body("1. Your invoices process themselves.")
add_body(
    "A document comes in — PDF, email, WhatsApp photo. AI reads it, extracts the data, "
    "matches it to your records, and flags exceptions. You approve. Done. What used to "
    "take 2 hours now takes 30 seconds."
)

add_bold_body("2. Your customers get answers at 2am.")
add_body(
    "Not a \"we'll get back to you\" message. Actual answers — product details, order status, "
    "pricing — from an AI assistant trained on your business. No overtime. No missed enquiries."
)

add_bold_body("3. Your leads don't go cold.")
add_body(
    "When someone enquires on WhatsApp, they get an instant, personalised response. Not a template. "
    "A real conversation that qualifies them and books a call. Response time drops from 24 hours "
    "to instant."
)

add_bold_body("4. Your reports write themselves.")
add_body(
    "Daily sales summaries, inventory alerts, outstanding payment reminders — generated and sent "
    "to you before your morning coffee. No one sits there compiling spreadsheets anymore."
)

add_bold_body("5. Your compliance checks run on autopilot.")
add_body(
    "Document verification, regulatory checklists, data validation — tasks that follow rules "
    "are perfect for AI. It checks everything, every time, without fatigue."
)

add_body(
    "None of this requires you to understand what a \"large language model\" is. "
    "You just need the right setup."
)

# ==================== SECTION 2 ====================
doc.add_paragraph("")
add_img("company_workflows.png")
add_heading_colored(
    "The Real Question Isn't \"Should I Use AI?\" — It's \"What Should I Automate First?\"",
    1, RGBColor(99, 102, 241)
)

add_body(
    "Most small businesses have 3–5 workflows that eat up hours every week but follow "
    "predictable patterns. These are your automation goldmines:"
)

bullets = [
    "Data entry from documents, forms, or messages",
    "Customer enquiries that ask the same 20 questions",
    "Follow-ups that fall through the cracks",
    "Reporting that someone compiles manually",
    "Compliance checks against long checklists",
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph("")
add_body(
    "If any of these sound familiar, you don't need a $50K digital transformation project. "
    "You need someone to sit with you for 30 minutes, map your workflows, and show you "
    "exactly what can be automated — and what the impact would be."
)

# ==================== SECTION 3 ====================
doc.add_paragraph("")
add_heading_colored(
    "Why Most Small Businesses Haven't Automated Yet",
    1, RGBColor(139, 92, 246)
)

add_body("It's not the cost. AI automation for MSMEs is more affordable than most people think — often less than the cost of one part-time hire.")
add_body("It's the jargon barrier.")
add_body(
    "When every AI company leads with \"we leverage cutting-edge LLMs with retrieval-augmented "
    "generation pipelines,\" most business owners (rightly) tune out. They don't want a technology "
    "lecture. They want to know:"
)

questions = [
    "Will this save me time?",
    "Will this save me money?",
    "How fast can it be set up?",
]
for q in questions:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.font.italic = True
    run.font.bold = True

doc.add_paragraph("")
add_body("The answers, for most MSME workflows, are: Yes. Yes. Days, not months.")

# ==================== LITMUS TEST ====================
doc.add_paragraph("")
add_img("company_litmus.png")
add_heading_colored("A Simple Litmus Test", 1, RGBColor(5, 150, 105))

add_body("Ask yourself:")
p = doc.add_paragraph()
run = p.add_run(
    "\"Is there a task my team does every day that follows roughly the same steps, "
    "takes more than 30 minutes, and doesn't require creative judgment?\""
)
run.font.italic = True
run.font.size = Pt(12)
run.font.bold = True

doc.add_paragraph("")
add_body(
    "If yes — that's automatable. Today. Without hiring a developer, without learning Python, "
    "and without understanding what GPT stands for."
)

# ==================== CLOSING ====================
doc.add_paragraph("")
doc.add_paragraph("—" * 40)
doc.add_paragraph("")

p = doc.add_paragraph()
run = p.add_run(
    "The businesses that win with AI aren't the ones with the biggest budgets. "
    "They're the ones that start with one workflow, see the result, and build from there."
)
run.font.bold = True
run.font.size = Pt(12)

doc.add_paragraph("")
p = doc.add_paragraph()
run = p.add_run("Start small. Start now.")
run.font.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(99, 102, 241)

doc.add_paragraph("")
doc.add_paragraph("")

# ==================== CTA (subtle) ====================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(
    "JSK AI helps small businesses automate their workflows — no jargon, no long contracts, "
    "just practical solutions delivered in days.\n"
    "Book a free 30-minute discovery call: https://jskai.ai/#cta"
)
run.font.italic = True
run.font.color.rgb = RGBColor(99, 102, 241)
run.font.size = Pt(10)

doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("#AIAutomation  #SmallBusiness  #MSME  #DigitalTransformation  #NoJargon")
run.font.color.rgb = RGBColor(120, 120, 140)
run.font.size = Pt(10)

output_path = os.path.join(output_dir, "LinkedIn_Article_Company_JSKAI.docx")
doc.save(output_path)
print(f"\nDocument saved to: {output_path}")
